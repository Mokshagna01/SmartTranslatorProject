"""
English → Indian Languages Document Translator
Preserves paragraph structure from PDF
"""

import os
import docx
import pdfplumber
from src.translation.translate import translate_paragraph


class EnglishToIndianDocTranslator:

    # -------- READ DOCUMENT --------
    def read_document(self, file_path):

        ext = os.path.splitext(file_path)[1].lower()

        blocks = []

        # TXT
        if ext == ".txt":

            with open(file_path, "r", encoding="utf-8") as f:
                blocks = f.readlines()

        # DOCX
        elif ext == ".docx":

            doc = docx.Document(file_path)

            for para in doc.paragraphs:
                blocks.append(para.text + "\n")

        # PDF
        elif ext == ".pdf":

            with pdfplumber.open(file_path) as pdf:

                for page in pdf.pages:

                    text = page.extract_text()

                    if text:
                        lines = text.split("\n")

                        for line in lines:
                            blocks.append(line + "\n")

        else:
            raise ValueError("Unsupported file format")

        return blocks


    # -------- SAVE DOCUMENT --------
    def save_document(self, blocks, output_file):

        ext = os.path.splitext(output_file)[1].lower()

        if ext == ".txt":

            with open(output_file, "w", encoding="utf-8") as f:

                for block in blocks:
                    f.write(block)

        elif ext == ".docx":

            doc = docx.Document()

            for block in blocks:
                doc.add_paragraph(block)

            doc.save(output_file)

        else:
            raise ValueError("Output must be TXT or DOCX")


    # -------- TRANSLATE DOCUMENT --------
    def translate_document(self, input_file, target_lang):

        blocks = self.read_document(input_file)

        translated_blocks = []

        print("\nTranslating document...\n")

        for i, block in enumerate(blocks):

            text = block.strip()

            if text == "":
                translated_blocks.append("\n")
                continue

            print(f"Translating line {i+1}/{len(blocks)}")

            translated = translate_paragraph(
                text,
                source_lang="eng_Latn",
                target_lang=target_lang
            )

            translated_blocks.append(translated + "\n")

        return translated_blocks