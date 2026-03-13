import os
from PyPDF2 import PdfReader
from docx import Document
from src.pipeline.phase1_global_to_english import translate_to_english as translate_paragraph

class GlobalToEnglishDocTranslator:
    def __init__(self):
        pass

    # -------- Read document ----------
    def read_document(self, file_path):
        _, ext = os.path.splitext(file_path.lower())

        paragraphs = []
        if ext == ".pdf":
            reader = PdfReader(file_path)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    paragraphs.extend([p.strip() for p in text.split("\n") if p.strip()])
        elif ext == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                paragraphs = [line.strip() for line in f.readlines() if line.strip()]
        elif ext == ".docx":
            doc = Document(file_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
        else:
            raise ValueError("Unsupported document type. Use PDF, TXT, or DOCX.")
        return paragraphs

    # -------- Translate ----------
    def translate_document(self, file_path, source_lang_code):
        paragraphs = self.read_document(file_path)
        translated_paragraphs = []

        for paragraph in paragraphs:
            translated_text = translate_paragraph(paragraph, source_lang_code)
            translated_paragraphs.append(translated_text)

        return translated_paragraphs

    # -------- Save document ----------
    def save_translated_document(self, translated_paragraphs, output_file):
        _, ext = os.path.splitext(output_file.lower())

        if ext == ".txt":
            with open(output_file, "w", encoding="utf-8") as f:
                for para in translated_paragraphs:
                    f.write(para + "\n\n")
        elif ext == ".docx":
            doc = Document()
            for para in translated_paragraphs:
                doc.add_paragraph(para)
            doc.save(output_file)
        else:
            raise ValueError("Unsupported output format. Use TXT or DOCX.")