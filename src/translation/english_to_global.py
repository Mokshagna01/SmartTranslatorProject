# src/translation/english_to_global.py

import os
import re
from PyPDF2 import PdfReader
from docx import Document
from src.pipeline.phase1_global_pipeline import hybrid_translate


class EnglishToGlobalDocTranslator:

    def read_document(self, file_path):

        ext = os.path.splitext(file_path)[1].lower()
        paragraphs = []

        # -------- PDF --------
        if ext == ".pdf":

            reader = PdfReader(file_path)

            buffer = ""

            for page in reader.pages:

                text = page.extract_text()

                if not text:
                    continue

                lines = text.split("\n")

                for line in lines:

                    line = line.strip()

                    if not line:
                        continue

                    buffer += " " + line

                    # paragraph end detection
                    if re.search(r'[.!?]$', line):

                        paragraphs.append(buffer.strip())
                        buffer = ""

            if buffer:
                paragraphs.append(buffer.strip())

        # -------- TXT --------
        elif ext == ".txt":

            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()

            paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

        # -------- DOCX --------
        elif ext == ".docx":

            doc = Document(file_path)

            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())

        else:
            raise ValueError("Unsupported file type")

        return paragraphs


    def save_translated_document(self, paragraphs, output_file):

        ext = os.path.splitext(output_file)[1].lower()

        if ext == ".txt":

            with open(output_file, "w", encoding="utf-8") as f:

                for para in paragraphs:
                    f.write(para + "\n\n")

        elif ext == ".docx":

            doc = Document()

            for para in paragraphs:
                doc.add_paragraph(para)

            doc.save(output_file)

        else:
            raise ValueError("Unsupported output format")


    def translate_document(self, input_file, target_language):

        paragraphs = self.read_document(input_file)

        translated_paragraphs = []

        for para in paragraphs:

            translated = hybrid_translate(
                para,
                src_lang="eng_Latn",
                tgt_lang=target_language
            )

            translated_paragraphs.append(translated)

        return translated_paragraphs